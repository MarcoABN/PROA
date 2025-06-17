import subprocess
from docx import Document
from tempfile import NamedTemporaryFile
from datetime import datetime
import os
import platform

def safe_str(value):
    return "" if value in [None, 0, "0", "0.0"] else str(value)

def safe_get(d, key):
    return d.get(key) if d.get(key) not in [None, 0, "0", "0.0"] else ""

def gerar_pdf_procuracao(cliente):
    doc = Document("docs/procuracao01OR.docx")

    print ('Entrou no utils', cliente)

    placeholders = {
        "${nomecliente1}": safe_get(cliente, "nome"),
        "${enderecocliente}": safe_get(cliente, 'logradouro'),
        "${cep}": safe_get(cliente, "cep"),
        "${cidade}": safe_get(cliente, "cidade"),
        "${bairro}": safe_get(cliente, "bairro"),
        "${rg}": safe_get(cliente, "rg"),
        "${orgexpedidor}": safe_get(cliente, "orgEmissor"),
        "${cpfcliente1}": safe_get(cliente, "cpfcnpj"),
        "${email}": safe_get(cliente, "email"),
        "${celular}": safe_get(cliente, "celular"),

        "${nomecliente2}": safe_get(cliente, "nome"),
        "${cpfcliente2}": safe_get(cliente, "cpfcnpj"),
        
    }

    # Substituir nos parágrafos
    for p in doc.paragraphs:
        for key, value in placeholders.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Substituir nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Salvar documento temporário
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        doc.save(tmp_docx.name)
        tmp_docx_path = tmp_docx.name

    output_dir = os.path.dirname(tmp_docx_path)

    # Verifica o sistema operacional
    libreoffice_cmd = "soffice" if platform.system() == "Windows" else "libreoffice"

    # Executar conversão
    try:
        subprocess.run([
            libreoffice_cmd,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            tmp_docx_path
        ], check=True)
    except FileNotFoundError as e:
        raise RuntimeError(f"Erro: não foi possível encontrar '{libreoffice_cmd}'. Verifique se o LibreOffice está instalado e no PATH.") from e

    tmp_pdf_path = tmp_docx_path.replace(".docx", ".pdf")

    # Ler PDF gerado
    with open(tmp_pdf_path, "rb") as pdf_file:
        pdf_bytes = pdf_file.read()

    # Limpar arquivos temporários
    os.unlink(tmp_docx_path)
    os.unlink(tmp_pdf_path)
    print ('vai retornar o pdf')
    return pdf_bytes
